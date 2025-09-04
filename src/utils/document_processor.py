"""
DALI Legal AI - Document Processing Utilities
Handles various document formats and text extraction
"""

import os
import logging
from typing import Optional, Dict, List
from pathlib import Path
import tempfile

# Document processing libraries
try:
    import PyPDF2
    from docx import Document as DocxDocument
    import openpyxl
    import markdown
except ImportError as e:
    logging.warning(f"Some document processing libraries not available: {e}")

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Document processor for various file formats
    Extracts text content from legal documents
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel
        }
    
    def process_file(self, file_input) -> Optional[str]:
        """
        Process a file and extract text content
        
        Args:
            file_input: File path (str/Path) or Streamlit uploaded file
            
        Returns:
            Extracted text content or None if processing failed
        """
        try:
            # Handle Streamlit uploaded file
            if hasattr(file_input, 'name') and hasattr(file_input, 'read'):
                return self._process_uploaded_file(file_input)
            
            # Handle file path
            file_path = Path(file_input)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            processor = self.supported_formats[file_extension]
            return processor(file_path)
            
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return None
    
    def _process_uploaded_file(self, uploaded_file) -> Optional[str]:
        """Process Streamlit uploaded file"""
        try:
            # Get file extension
            file_extension = Path(uploaded_file.name).suffix.lower()
            
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}")
                return None
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                tmp_file.write(uploaded_file.read())
                tmp_file_path = tmp_file.name
            
            try:
                # Process the temporary file
                processor = self.supported_formats[file_extension]
                content = processor(Path(tmp_file_path))
                return content
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"Error processing uploaded file: {e}")
            return None
    
    def _process_pdf(self, file_path: Path) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content.append(page.extract_text())
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return None
    
    def _process_docx(self, file_path: Path) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return None
    
    def _process_doc(self, file_path: Path) -> Optional[str]:
        """Extract text from DOC file (legacy format)"""
        try:
            # For .doc files, we'd need python-docx2txt or similar
            # For now, return a message suggesting conversion
            logger.warning(f"Legacy DOC format not fully supported: {file_path}")
            return f"Legacy DOC file detected: {file_path.name}. Please convert to DOCX format for better processing."
            
        except Exception as e:
            logger.error(f"Error processing DOC {file_path}: {e}")
            return None
    
    def _process_txt(self, file_path: Path) -> Optional[str]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
                
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            logger.error(f"Could not decode text file: {file_path}")
            return None
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            return None
    
    def _process_markdown(self, file_path: Path) -> Optional[str]:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to plain text (remove formatting)
            html = markdown.markdown(md_content)
            
            # Basic HTML tag removal (for better text extraction)
            import re
            text = re.sub(r'<[^>]+>', '', html)
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text
            
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
            return None
    
    def _process_excel(self, file_path: Path) -> Optional[str]:
        """Extract text from Excel file"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_text.append(str(cell_value))
                    
                    if row_text:
                        text_content.append(' | '.join(row_text))
                
                text_content.append('')  # Empty line between sheets
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            return None
    
    def extract_metadata(self, file_path: Path) -> Dict:
        """Extract metadata from document"""
        try:
            stat = file_path.stat()
            
            metadata = {
                'filename': file_path.name,
                'file_size': stat.st_size,
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime,
                'file_extension': file_path.suffix.lower()
            }
            
            # Format-specific metadata
            if file_path.suffix.lower() == '.pdf':
                metadata.update(self._extract_pdf_metadata(file_path))
            elif file_path.suffix.lower() == '.docx':
                metadata.update(self._extract_docx_metadata(file_path))
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata from {file_path}: {e}")
            return {}
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict:
        """Extract PDF-specific metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'page_count': len(pdf_reader.pages),
                    'pdf_version': pdf_reader.pdf_header if hasattr(pdf_reader, 'pdf_header') else None
                }
                
                # Extract document info if available
                if pdf_reader.metadata:
                    doc_info = pdf_reader.metadata
                    metadata.update({
                        'title': doc_info.get('/Title', ''),
                        'author': doc_info.get('/Author', ''),
                        'subject': doc_info.get('/Subject', ''),
                        'creator': doc_info.get('/Creator', ''),
                        'producer': doc_info.get('/Producer', ''),
                        'creation_date': doc_info.get('/CreationDate', ''),
                        'modification_date': doc_info.get('/ModDate', '')
                    })
                
                return metadata
                
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {e}")
            return {}
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict:
        """Extract DOCX-specific metadata"""
        try:
            doc = DocxDocument(file_path)
            
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'table_count': len(doc.tables)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'keywords': core_props.keywords or '',
                    'comments': core_props.comments or '',
                    'created': core_props.created.isoformat() if core_props.created else '',
                    'modified': core_props.modified.isoformat() if core_props.modified else ''
                })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {e}")
            return {}
    
    def validate_document(self, file_path: Path) -> Dict:
        """Validate document and return status"""
        try:
            if not file_path.exists():
                return {'valid': False, 'error': 'File does not exist'}
            
            if file_path.stat().st_size == 0:
                return {'valid': False, 'error': 'File is empty'}
            
            file_extension = file_path.suffix.lower()
            if file_extension not in self.supported_formats:
                return {'valid': False, 'error': f'Unsupported format: {file_extension}'}
            
            # Try to process a small portion to validate
            try:
                content = self.process_file(file_path)
                if content and len(content.strip()) > 0:
                    return {
                        'valid': True,
                        'content_length': len(content),
                        'preview': content[:200] + '...' if len(content) > 200 else content
                    }
                else:
                    return {'valid': False, 'error': 'No text content could be extracted'}
                    
            except Exception as e:
                return {'valid': False, 'error': f'Processing error: {str(e)}'}
                
        except Exception as e:
            return {'valid': False, 'error': f'Validation error: {str(e)}'}
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.supported_formats.keys())
    
    def estimate_processing_time(self, file_path: Path) -> float:
        """Estimate processing time based on file size and type"""
        try:
            file_size = file_path.stat().st_size
            file_extension = file_path.suffix.lower()
            
            # Base processing time per MB
            time_per_mb = {
                '.txt': 0.1,
                '.md': 0.1,
                '.docx': 0.5,
                '.pdf': 1.0,
                '.xlsx': 0.8,
                '.doc': 1.5
            }
            
            base_time = time_per_mb.get(file_extension, 1.0)
            file_size_mb = file_size / (1024 * 1024)
            
            return max(0.5, base_time * file_size_mb)  # Minimum 0.5 seconds
            
        except Exception:
            return 5.0  # Default estimate


# Utility functions
def is_text_file(file_path: Path) -> bool:
    """Check if file is a text-based document"""
    text_extensions = {'.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm'}
    return file_path.suffix.lower() in text_extensions


def get_file_type_description(file_extension: str) -> str:
    """Get human-readable description of file type"""
    descriptions = {
        '.pdf': 'PDF Document',
        '.docx': 'Microsoft Word Document',
        '.doc': 'Legacy Microsoft Word Document',
        '.txt': 'Plain Text File',
        '.md': 'Markdown Document',
        '.xlsx': 'Microsoft Excel Spreadsheet',
        '.xls': 'Legacy Microsoft Excel Spreadsheet'
    }
    
    return descriptions.get(file_extension.lower(), f'{file_extension.upper()} File')


if __name__ == "__main__":
    # Example usage
    processor = DocumentProcessor()
    
    print("Supported formats:", processor.get_supported_formats())
    
    # Test with a sample file (if available)
    test_file = Path("sample.txt")
    if test_file.exists():
        content = processor.process_file(test_file)
        if content:
            print(f"Extracted {len(content)} characters from {test_file}")
            print(f"Preview: {content[:100]}...")
        else:
            print("Failed to extract content")
    else:
        print("No test file available")

